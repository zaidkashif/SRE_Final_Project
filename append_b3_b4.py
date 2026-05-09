from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def append_part_b3_b4():
    doc = Document('SRE_Final_Project_Report.docx')
    
    # Part B3
    doc.add_page_break()
    doc.add_heading('B3. Smell Interaction and Prioritisation [3 Marks]', level=2)
    
    p = doc.add_paragraph()
    p.add_run('1. Interaction of Smells (Long Method & Data Clumps):\n').bold = True
    p.add_run('In HelpFormatter.java, the "Long Method" smell in appendOptions() interacts heavily with the "Data Clumps" smell. The method takes isolated primitives (width, leftPad, descPad) which dictate text wrapping context. Because these variables are not bound into an object, they are passed as raw arguments through chained methods (appendWrappedText, renderWrappedText, printWrapped). This forces appendOptions to act as a massive procedural coordinating block that calculates shifts in padding, expanding its length and violating the Single Responsibility Principle. The lack of an abstraction (Data Clump) directly inflates the procedural footprint (Long Method).\n\n')
    
    p.add_run('2. Greatest Risk to Maintainability:\n').bold = True
    p.add_run('The greatest risk to long-term maintainability is the "Divergent Change" and "Feature Envy" present in the Option and HelpFormatter classes. Specifically, HelpFormatter is tightly coupled to the Option model. If a developer needs to introduce a new command-line constraint or style (e.g., colored text or nested option groups), they face tracing through the 80-line appendOptions() method. The risk of breaking text alignment or infinite loops (as seen in the while loop of appendWrappedText) makes future modifications highly error-prone.\n\n')
    
    p.add_run('3. Priority Refactoring Treatment:\n').bold = True
    p.add_run('The very first treatment applied should be "Introduce Parameter Object" to cure the Data Clumps. Consolidating the width, leftPad, descPad, and nextLineTabStop into a TextFormattingContext object requires minimal structural logic shifts (high safety) but delivers massive readability improvements, paving the way to later perform "Extract Method" on appendOptions() with much cleaner signatures.')

    # Part B4
    doc.add_heading('B4. Refactoring Demonstration [3 Marks]', level=2)
    
    p_orig = doc.add_paragraph()
    p_orig.add_run('Original Code Snippet (Smelly Version - Data Clumps):\n').bold = True
    p_orig.add_run('Location: HelpFormatter.java (Lines 909-923)\n').italic = True
    doc.add_paragraph('public void printWrapped(final PrintWriter pw, final int width, final int nextLineTabStop, final String text) {\n    pw.println(renderWrappedTextBlock(new StringBuilder(text.length()), width, nextLineTabStop, text));\n}\n\npublic void printWrapped(final PrintWriter pw, final int width, final String text) {\n    printWrapped(pw, width, 0, text);\n}')
    
    p_ref = doc.add_paragraph()
    p_ref.add_run('Refactored Code Snippet (Improved Version - Introduce Parameter Object):\n').bold = True
    doc.add_paragraph('// New Parameter Object extracted\npublic class TextWrapContext {\n    public final int width;\n    public final int nextLineTabStop;\n    \n    public TextWrapContext(int width, int nextLineTabStop) {\n        this.width = width;\n        this.nextLineTabStop = nextLineTabStop;\n    }\n}\n\n// Cleaned up method signatures\npublic void printWrapped(final PrintWriter pw, final TextWrapContext context, final String text) {\n    pw.println(renderWrappedTextBlock(new StringBuilder(text.length()), context.width, context.nextLineTabStop, text));\n}\n\npublic void printWrapped(final PrintWriter pw, final int width, final String text) {\n    printWrapped(pw, new TextWrapContext(width, 0), text);\n}')
    
    p_eval = doc.add_paragraph()
    p_eval.add_run('Behavior Preservation & Structural Improvement:\n').bold = True
    p_eval.add_run('The external behavior is wholly unchanged because the default configuration dynamically generates the necessary wrapper with default tab stops (0). The caller relying on the simpler signature notices no difference. Structurally, we have removed the disconnected data parameters from the print/render pipeline, setting up a solid foundation to add additional formatting properties (like indent styles arrays or prefixes) directly into TextWrapContext without continuously breaking method chain signatures.')

    doc.save('SRE_Final_Project_Report.docx')

if __name__ == '__main__':
    append_part_b3_b4()
