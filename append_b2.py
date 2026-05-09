from docx import Document

def append_part_b2():
    doc = Document('SRE_Final_Project_Report.docx')
    
    doc.add_page_break()
    doc.add_heading('B2. The Five Code Smell Categories — Deep Identification [16 Marks]', level=2)
    doc.add_paragraph('All 5 smell categories have been observed in the apache/commons-cli codebase.')

    # 1. Bloaters
    doc.add_heading('Category 1: Bloaters [4 Marks]', level=3)
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Smell Name'
    hdr[1].text = 'File and Line Number'
    hdr[2].text = 'Evidence — what you see in the code'
    hdr[3].text = 'Why It Is a Bloater / Recommended Treatment'
    
    row1 = table1.add_row().cells
    row1[0].text = 'Long Method'
    row1[1].text = 'HelpFormatter.java : 423'
    row1[2].text = 'appendOptions(final A sb, final int width, final Options options...)'
    row1[3].text = 'This method is nearly 80 lines long. It handles word wrapping, string appending, calculating padding widths across all options, checking deprecation formats, and rendering output.\n\nTreatment: Extract Method (e.g., extract the width-calculation logic into a calculateOptimalWidth method).'
    
    row2 = table1.add_row().cells
    row2[0].text = 'Data Clumps'
    row2[1].text = 'HelpFormatter.java : 909'
    row2[2].text = 'The parameters (int width, int nextLineTabStop, String text) repeatedly appear together.'
    row2[3].text = 'These three primitives are passed together across printWrapped, renderWrappedText, and appendWrappedText, acting as a unit without being grouped.\n\nTreatment: Introduce Parameter Object (e.g., TextWrapContext).'

    p1 = doc.add_paragraph('\n')
    p1.add_run('Code Snippet Evidence (Long Method - HelpFormatter.java):\n').bold = True
    p1.add_run('<A extends Appendable> A appendOptions(final A sb, final int width, final Options options, final int leftPad, final int descPad) throws IOException {\n    final String lpad = createPadding(leftPad);\n    final String dpad = createPadding(descPad);\n    // ... 75 more lines of logic ...\n}')

    # 2. OO-Abusers
    doc.add_heading('Category 2: Object-Orientation Abusers [3 Marks]', level=3)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Smell Name'
    hdr[1].text = 'File and Line Number'
    hdr[2].text = 'Evidence from Your Code'
    hdr[3].text = 'OO Principle Violated / Treatment'
    
    row1 = table2.add_row().cells
    row1[0].text = 'Switch Statements'
    row1[1].text = 'PatternOptionBuilder.java : 133'
    row1[2].text = 'switch (ch) { case \'@\': return OBJECT_VALUE; case \':\': return STRING_VALUE; ... }'
    row1[3].text = 'Violates Open/Closed Principle. Hardcoded character-to-class mapping inside a switch forces modifying this method each time a new type is supported.\n\nTreatment: Replace Conditional with Polymorphism (or Strategy Pattern via a registered Map).'
    
    row2 = table2.add_row().cells
    row2[0].text = 'Temporary Fields'
    row2[1].text = 'DefaultParser.java : 204'
    row2[2].text = 'protected CommandLine cmd;\nprotected Option currentOption;'
    row2[3].text = 'Violates encapsulation. These fields are defined at the class level but are exclusively meaningful during the brief execution of the parse() method.\n\nTreatment: Extract Class (e.g., parsing state could be moved into a ParserState object passed through contexts).'
    
    # 3. Change Preventors
    doc.add_heading('Category 3: Change Preventors [3 Marks]', level=3)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Smell Name'
    hdr[1].text = 'File(s) and Lines'
    hdr[2].text = 'How Many Places Must Change? / Treatment Strategy'
    
    row1 = table3.add_row().cells
    row1[0].text = 'Divergent Change'
    row1[1].text = 'HelpFormatter.java : 42'
    row1[2].text = '1 class must change for formatting, wrapping algorithms, and output streams.\n\nTreatment Strategy: Extract Class (separate logical String algorithms from IO formatting).'
    
    row2 = table3.add_row().cells
    row2[0].text = 'Shotgun Surgery'
    row2[1].text = 'Option.java : 60'
    row2[2].text = 'Changing a core Option identity (e.g., adding `since` versioning) touched Option, OptionBuilder, Options, HelpFormatter, and CommandLine internally.\n\nTreatment Strategy: Move Field/Method to abstract away direct property reads from consumers.'

    # 4. Dispensables
    doc.add_heading('Category 4: Dispensables [3 Marks]', level=3)
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Smell Name'
    hdr[1].text = 'File and Line Number'
    hdr[2].text = 'What Makes It Dispensable? / Treatment'
    
    row1 = table4.add_row().cells
    row1[0].text = 'Lazy Class'
    row1[1].text = 'BasicParser.java : 26'
    row1[2].text = 'The class only overrides one method flatten() to return the arguments array untouched. It has almost no behavior of its own.\n\nTreatment: Inline Class or Collapse Hierarchy.'
    
    row2 = table4.add_row().cells
    row2[0].text = 'Comments (Redundant)'
    row2[1].text = 'OptionBuilder.java : 28'
    row2[2].text = '`/** Long option. */ private static String longOption;` The comment is completely redundant and adds zero semantic value over the variable name.\n\nTreatment: Delete the comment.'

    # 5. Couplers
    doc.add_heading('Category 5: Couplers [3 Marks]', level=3)
    table5 = doc.add_table(rows=1, cols=4)
    table5.style = 'Table Grid'
    hdr = table5.rows[0].cells
    hdr[0].text = 'Smell Name'
    hdr[1].text = 'File and Line Number'
    hdr[2].text = 'Description of the Coupling Problem'
    hdr[3].text = 'Treatment'
    
    row1 = table5.add_row().cells
    row1[0].text = 'Feature Envy'
    row1[1].text = 'HelpFormatter.java : 442'
    row1[2].text = 'The HelpFormatter constantly calls Option getters (hasLongOpt, getArgName, isDeprecated) to determine how to format strings.'
    row1[3].text = 'Move Method. Option should probably expose a local getFormattedSignature() method so HelpFormatter doesn\'t need to poke at its internals.'
    
    row2 = table5.add_row().cells
    row2[0].text = 'Message Chain'
    row2[1].text = 'DefaultParser.java : 383'
    row2[2].text = '`options.getOption(token).getLongOpt()` - It calls the target, gets an object back, and immediately asks the returned object for its field.'
    row2[3].text = 'Hide Delegate. Create a getLongOptForPrefix(token) method directly in the Options object.'

    doc.save('SRE_Final_Project_Report.docx')

if __name__ == '__main__':
    append_part_b2()
